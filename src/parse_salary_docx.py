from docx import Document
import re

def parse_salary_docx(docx_path: str) -> dict:
    """
    Parse salary details and employment terms from DOCX file.
    """
    try:
        doc = Document(docx_path)
        all_text = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            all_text.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if cells:
                    all_text.append(" ".join(cells))

        # Clean text
        text = "\n".join(all_text).replace("\t", " ").replace(",", "").lower()

        # Enhanced patterns for salary and employment terms
        patterns = {
            # Salary Components
            "Basic": r"basic\s*salary\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Housing": r"housing\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Transport": r"transport\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Travel": r"travel\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Accommodation": r"accommodation\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Plane": r"plane\s*:?\s*aed\s*([\d,]+\.?\d*)",
            "Overtime": r"overtime\s*\+\s*premium\s*aed\s*([\d,]+\.?\d*)",
            "Site_Allowance": r"site\s*allowance\s*aed\s*([\d,]+\.?\d*)",
            "Total": r"total\s*monthly\s*package\s*:?\s*aed\s*([\d,]+\.?\d*)",
            
            # Employment Terms
            "Probation_Period": r"probation\s*period\s*:?\s*(\d+\s*(?:months?|days?|weeks?))",
            "Notice_Period": r"notice\s*period\s*:?\s*(\d+\s*(?:months?|days?|weeks?))",
            "Working_Hours": r"(?:working\s*hours?|hours?\s*of\s*work)\s*:?\s*(\d+(?:\s*-\s*\d+)?\s*(?:hours?|hrs?))",
            "Contract_Duration": r"contract\s*duration\s*:?\s*(\d+\s*(?:months?|years?))",
            "Annual_Leave": r"annual\s*leave\s*:?\s*(\d+\s*(?:days?|weeks?))",
            "Sick_Leave": r"sick\s*leave\s*:?\s*(\d+\s*(?:days?|weeks?))",
            "Work_Days": r"work\s*days?\s*:?\s*(\d+\s*days?\s*per\s*week)",
            "Overtime_Rate": r"overtime\s*rate\s*:?\s*(\d+(?:\.\d+)?\s*(?:times?|x))",
            "Gratuity": r"gratuity\s*:?\s*(\d+(?:\.\d+)?\s*(?:months?|days?))",
            "Insurance": r"(?:health\s*)?insurance\s*:?\s*(provided|included|yes|no)",
            "Visa_Status": r"visa\s*status\s*:?\s*(provided|included|yes|no|sponsored)",
            "Flight_Ticket": r"flight\s*ticket\s*:?\s*(provided|included|yes|no|annual|biannual)",
            "Accommodation_Provided": r"accommodation\s*:?\s*(provided|included|yes|no|shared|single)",
            "Transport_Provided": r"transport\s*:?\s*(provided|included|yes|no|shared|private)"
        }

        # Extract all values
        extracted_data = {}
        for key, pattern in patterns.items():
            match = re.search(pattern, text)
            if match:
                value = match.group(1)
                if key in ["Insurance", "Visa_Status", "Flight_Ticket", "Accommodation_Provided", "Transport_Provided"]:
                    # For boolean/status fields, keep the original text
                    extracted_data[key] = value.title()
                else:
                    # For numeric fields, add AED if it's a salary component
                    if key in ["Basic", "Housing", "Transport", "Travel", "Accommodation", "Plane", "Overtime", "Site_Allowance", "Total"]:
                        extracted_data[key] = f"{value} AED"
                    else:
                        extracted_data[key] = value

        # Process allowances logic (same as before)
        exclude_from_combination = ["Travel", "Accommodation", "Plane", "Transport"]
        
        combinable_allowances = []
        for key in ["Overtime", "Site_Allowance"]:
            if key in extracted_data:
                combinable_allowances.append((key, extracted_data[key]))

        # Build final salary data
        salary_data = {}
        
        # Add basic and total
        if "Basic" in extracted_data:
            salary_data["Basic"] = extracted_data["Basic"]
        if "Total" in extracted_data:
            salary_data["Total"] = extracted_data["Total"]
            
        # Add excluded allowances individually
        for key in exclude_from_combination:
            if key in extracted_data:
                salary_data[key] = extracted_data[key]

        # Process combinable allowances
        if len(combinable_allowances) == 0:
            pass
        elif len(combinable_allowances) == 1:
            key, value = combinable_allowances[0]
            if key == "Overtime":
                salary_data["Overtime_Allowance"] = value
            elif key == "Site_Allowance":
                salary_data["Site_Allowance"] = value
            else:
                salary_data["Other_Allowance"] = value
        else:
            total_amount = 0
            allowance_names = []
            
            for key, value in combinable_allowances:
                amount_str = value.replace(" AED", "").replace(",", "")
                try:
                    amount = float(amount_str)
                    total_amount += amount
                    allowance_names.append(key)
                except ValueError:
                    continue
            
            if total_amount > 0:
                salary_data["Other_Allowances"] = f"{total_amount:,.0f} AED"
                breakdown = " + ".join([f"{name}" for name in allowance_names])
                salary_data["Other_Allowances_Breakdown"] = breakdown

        # Add employment terms
        employment_terms = {}
        for key in ["Probation_Period", "Notice_Period", "Working_Hours", "Contract_Duration", 
                   "Annual_Leave", "Sick_Leave", "Work_Days", "Overtime_Rate", "Gratuity",
                   "Insurance", "Visa_Status", "Flight_Ticket", "Accommodation_Provided", "Transport_Provided"]:
            if key in extracted_data:
                employment_terms[key] = extracted_data[key]

        # Combine salary and employment terms
        if employment_terms:
            salary_data["Employment_Terms"] = employment_terms

        return salary_data

    except Exception as e:
        print(f"⚠️ Failed to parse salary DOCX: {e}")
        return {}
